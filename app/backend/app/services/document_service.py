import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.document import Document
from app.models.return_request import ReturnRequest


# Документы, которые формирует сама АИС (внутренние, не учётные)
DOCUMENT_TYPES = {
    "claim_letter": "Претензионное письмо заводу",
    "route_sheet": "Маршрутный лист",
    "inspection_act": "Акт осмотра / сверки товара",
    "rejection_notice": "Уведомление об отказе в возврате",
}

# Учётные документы, которые формируются в 1С и подтягиваются в АИС
# (АИС их не генерирует, а сохраняет № и печатную форму, полученные из 1С)
ONEC_DOCUMENT_TYPES = {
    "write_off": "Акт списания товара (1С)",
    "correction": "Корректировка / возврат в продажу (1С)",
    "reconciliation_act": "Акт сверки (1С)",
}


def _ensure_dir(path: str):
    Path(path).mkdir(parents=True, exist_ok=True)


def generate_docx(return_request: ReturnRequest, doc_type: str) -> tuple[str, str]:
    """Generate a .docx document and return (file_name, file_path)."""
    _ensure_dir(settings.DOCUMENTS_DIR)

    title = DOCUMENT_TYPES.get(doc_type, doc_type)
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    file_name = f"{doc_type}_{return_request.number}_{timestamp}.docx"
    file_path = os.path.join(settings.DOCUMENTS_DIR, file_name)

    doc = DocxDocument()

    # Header
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Заявка: {return_request.number}")
    doc.add_paragraph(
        f"Дата формирования: {datetime.now(timezone.utc).strftime('%d.%m.%Y %H:%M')}"
    )
    doc.add_paragraph("")

    # Client info
    if return_request.client:
        doc.add_heading("Сведения о покупателе", level=2)
        doc.add_paragraph(f"ФИО / Организация: {return_request.client.name}")
        if return_request.client.phone:
            doc.add_paragraph(f"Телефон: {return_request.client.phone}")
        if return_request.client.email:
            doc.add_paragraph(f"Email: {return_request.client.email}")

    # Return info
    doc.add_heading("Сведения о возврате", level=2)
    doc.add_paragraph(f"Тип возврата: {return_request.return_type}")
    if return_request.reason:
        doc.add_paragraph(f"Причина: {return_request.reason.name}")
    if return_request.comment:
        doc.add_paragraph(f"Комментарий: {return_request.comment}")

    # Items table
    if return_request.items:
        doc.add_heading("Товарные позиции", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Наименование", "Артикул", "Кол-во", "Ед.", "Цена, руб."]
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        for item in return_request.items:
            row = table.add_row()
            row.cells[0].text = item.product_name
            row.cells[1].text = item.article or ""
            row.cells[2].text = str(item.quantity)
            row.cells[3].text = item.unit
            row.cells[4].text = f"{item.price:.2f}"

        doc.add_paragraph(f"\nИтого: {return_request.total_amount:.2f} руб.")

    # Signatures
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Подпись менеджера                    Подпись покупателя")

    doc.save(file_path)
    return file_name, file_path


async def generate_and_save_document(
    db: AsyncSession, return_request: ReturnRequest, doc_type: str
) -> Document:
    """Generate document and save metadata to DB."""
    file_name, file_path = generate_docx(return_request, doc_type)

    document = Document(
        return_request_id=return_request.id,
        document_type=doc_type,
        file_name=file_name,
        file_path=file_path,
        generated_by="system",
    )
    db.add(document)
    await db.flush()
    return document


def register_onec_document(
    db_session, return_request_id: int, onec_doc_type: str,
    onec_number: str, file_path: str | None = None
) -> Document:
    """Зарегистрировать в АИС учётный документ, сформированный в 1С.

    АИС не генерирует этот документ, а сохраняет его номер и (при наличии)
    путь к печатной форме, полученной из 1С, привязывая к заявке.
    Вызывается из синхронной задачи обмена (Celery), поэтому работает с
    обычной (sync) сессией.
    """
    title = ONEC_DOCUMENT_TYPES.get(onec_doc_type, onec_doc_type)
    document = Document(
        return_request_id=return_request_id,
        document_type=onec_doc_type,
        file_name=f"{title} {onec_number}",
        file_path=file_path or "",
        generated_by="1С",
    )
    db_session.add(document)
    db_session.flush()
    return document
